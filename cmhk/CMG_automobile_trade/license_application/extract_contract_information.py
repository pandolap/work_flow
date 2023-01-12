#!/usr/bin/env python
# -*- coding: UTF-8 -*-
import datetime
import pandas as pd
import os
import re
from docx import Document
import win32com.client as win32

# 入参
config_dict = {
    'base_path': 'D:\RPA\许可证'
}

# start
contract_path = os.path.join(config_dict.get('base_path'), '合同')
# 月份处理字典
month_dict = {
    'JAN': '01',
    'FEB': '02',
    'MAR': '03',
    'APR': '04',
    'MAY': '05',
    'JUN': '06',
    'JUL': '07',
    'AUG': '08',
    'SEP': '09',
    'OCT': '10',
    'NOV': '11',
    'DEC': '12'
}


def get_file_path_list(folder):
    """
    获取目标文件夹中所有文件地址列表
    :param folder: 目标文件夹
    :return: 文件地址列表
    """
    if folder is None or folder == '':
        raise Exception('业务异常-获取合同数据：检测合同文件夹为空')

    file_names = os.listdir(folder)
    contract_file_list = []
    for file_name in file_names:
        contract_file_list.append(os.path.join(folder, file_name))
    return contract_file_list


def doc_to_docx(word_path):
    """
    简单的将doc转换为docx文档
    :param word_path: doc文档地址
    :return: to_path: 转换后的docx文档地址
    """
    word = win32.Dispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(word_path)
    to_path = "{}x".format(word_path)
    if os.path.exists(to_path):
        os.remove(to_path)
    doc.SaveAs(to_path, 12)
    doc.Close()
    word.Application.Quit()
    # 删除旧文件
    if os.path.exists(to_path):
        os.remove(word_path)
    return to_path


def parse_word(file_path):
    """
    解析docx文件获取解析到的合同数据
    :param file_path:
    :return:
    """
    if not file_path or not os.path.exists(file_path):
        raise Exception('业务异常-解析docx文件数据：%s文件按路径异常' % file_path)

    doc = Document(file_path)

    # 获取签订日期以及合同编号
    paragraph = doc.paragraphs
    # 过滤其中的空行
    paragraphs = list(filter(None, paragraph))
    # 获取第二行段落
    to_text = paragraphs[1].text
    # 判断文本中是中文符号还是英文符号进行分割
    if to_text.find(':') == -1:
        process_set = to_text.replace('：', '').split(' ')
    else:
        process_set = to_text.replace(':', '').split(' ')
    process_set = list(filter(None, process_set))
    # 获取外贸合同号
    flag = False
    contract_no = ''
    for process in process_set:
        if flag:
            contract_no = process
            break
        if process == 'No.':
            flag = True
    # 检查是否获取到合同号
    if not contract_no:
        contract_no = process_set[2]
    # 获取日期
    date_process = process_set[-2:]
    # 获取月份
    f = date_process[0].split(',')
    if len(f) == 2:
        f = f[1]
    else:
        f = f[0]
    _month = month_dict.get(f.upper(), '')
    # 获取绝体天数
    _year = date_process[-1][-4:]
    day_str = date_process[-1][:-4]
    _day = re.findall(r"\d+", day_str)[-1]
    try:
        sheet_data_fmt = datetime.date(int(_year), int(_month), int(_day)).strftime('%Y-%m-%d')
    except ValueError as e:
        raise Exception('程序异常-读取合同信息：读取到合同签订日期异常“%s”,异常信息：%s' % (date_process[-1], str(e)))
    # 获取规格型号
    tables = doc.tables
    table_data = []
    for table in tables:
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                row_content.append(cell.text)
            table_data.append(row_content)

    total_amount_process = table_data[-1]
    total_amount = total_amount_process[1].split('\n')[0].split(' ')[-1]

    # 规格数据
    specs = table_data[2:-1]

    entity = {
        "合同单号": contract_no,
        "签订日期": sheet_data_fmt,
        "总金额": total_amount,
        "规格型号": specs
    }

    return entity


def read_file_data(file_path):
    """
    读取传入的文件地址，读取文件中的数据并返回
    :param file_path: 文件链接
    :return: 获取的数据实体
    """
    # 将读取文件拆分为word 和 excel两种方法
    if file_path is None or not os.path.exists(file_path):
        raise Exception('业务异常-获取合同数据：%s文件路径异常' % file_path)

    df = pd.read_excel(file_path)
    sheet_data = df.values
    # 记录关键点的位置信息
    index_list = []
    for idx, row in df.iterrows():
        for i, col in enumerate(row):
            if 'Contract No.：' == col:
                index_list.append({"外贸合同号": {'r': idx, 'c': i}})
            if 'Signed Place, Date：' == col:
                index_list.append({"签订日期": {'r': idx, 'c': i}})
            if 'TOTAL AMOUNT:' == col:
                index_list.append({"总金额": {'r': idx, 'c': i}})
            if 'Commodity and Specifications' == col:
                index_list.append({"规格列表": {'r': idx, 'c': i}})

    # 获取合同号
    contract_r = index_list[0].get('外贸合同号').get('r')
    contract_c = index_list[0].get('外贸合同号').get('c')
    while True:
        contract_c += 1
        contract_no = sheet_data[contract_r][contract_c]
        if not pd.isna(contract_no):
            break

    # 获取签订日期
    # 合同文件有概率出现年份与日期粘连的情况
    signing_r = index_list[1].get('签订日期').get('r')
    signing_c = index_list[1].get('签订日期').get('c')
    while True:
        signing_c += 1
        signing_date_str = sheet_data[signing_r][signing_c]
        if not pd.isna(signing_date_str):
            break

    temp_list = signing_date_str.replace(',', ' ').split(" ")[-3:]
    # temp_list = signing_date_str.replace(',', '').split(' ')[1:]
    # if len(temp_list) != 2:
    #     raise Exception('程序异常-提取合同信息：获取合同的签订日期解析异常“%s”' % signing_date_str)
    year_str = temp_list[-1][-4:]
    month_str = month_dict.get(temp_list[0].upper())
    # day_str_ = temp_list[-1][:-4]
    day_str_ = temp_list[-2]
    day_str = re.findall(r"\d+", day_str_)[-1]
    try:
        sheet_data_fmt = datetime.date(int(year_str), int(month_str), int(day_str)).strftime('%Y-%m-%d')
    except ValueError as e:
        raise Exception('程序异常-读取合同信息：读取到合同签订日期异常“%s”,异常信息：%s' % (signing_date_str, str(e)))
    # 获取总金额
    total_amount = 0
    total_r = index_list[3].get('总金额').get('r')
    for col in sheet_data[total_r]:
        if not pd.isna(col) and col != 'TOTAL AMOUNT:':
            total_amount_text = col.strip().split(' ')[-1].replace(',', '')
            total_amount = '.'.join(re.findall(r"\d+", total_amount_text))
            break

    # 商品规格集合
    super_index = index_list[2].get('规格列表').get('r') + 2
    sub_index = index_list[3].get('总金额').get('r')
    specs = []
    ca = []
    for i in range(super_index, sub_index):
        if len(ca) > 0:
            ca = []
        for col in sheet_data[i]:
            if not pd.isna(col):
                ca.append(col)
        specs.append(ca)

    entity = {
        "合同单号": contract_no,
        "签订日期": sheet_data_fmt,
        "总金额": total_amount,
        "规格型号": specs
    }

    return entity


contract_list = get_file_path_list(contract_path)
contract_data_list = []
for contract in contract_list:
    (file, ext) = os.path.splitext(contract)
    if ext == '.docx':
        enti = parse_word(contract)
        contract_data_list.append(enti)
    elif ext == '.xls' or ext == '.xlsx':
        enti = read_file_data(contract)
        contract_data_list.append(enti)
    elif ext == '.doc':
        out_file = doc_to_docx(contract)
        enti = parse_word(out_file)
        contract_data_list.append(enti)
print('合同解析数据：' + str(contract_data_list))
# end
